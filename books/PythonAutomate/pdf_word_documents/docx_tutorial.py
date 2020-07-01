"""docx_tutorial.py
튜토리얼 코드 : https://python-docx.readthedocs.io/en/latest/
"""
from docx import Document
from docx.shared import Inches

document = Document()

# 헤더 추가 -> H0
document.add_heading("Document Title", 0)

# paragraph, run 추가 및 속성 적용
p = document.add_paragraph("A plain paragraph having some ")
p.add_run("bold").bold = True
p.add_run(" and some ")
p.add_run("italic").italic = True

document.add_heading("Heading, level 1", level=1)
# 스타일 적용
document.add_paragraph("Intense quote", style="Intense Quote")
document.add_paragraph("first item in unordered list", style="List Bullet")
document.add_paragraph("first item in ordered list", style="List Number")

# 사진 추가하기
document.add_picture("catlogo.png", width=Inches(1.25))

records = (
    (3, "101", "Spam"),
    (7, "422", "Eggs"),
    (4, "631", "Spam, spam, eggs, and spam")
)

# 표 추가하기
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Qty"
hdr_cells[1].text = "Id"
hdr_cells[2].text = "Desc"
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save("demo.docx")
