"""reading_word.py
워드 문서 읽기
"""
import docx

doc = docx.Document("demo.docx")

print(len(doc.paragraphs))  # paragraph 갯수
print(doc.paragraphs[0].text)  # 첫번째 paragraph 문자열
print(doc.paragraphs[1].text)  # 두번째 paragraph 문자열

print(len(doc.paragraphs[1].runs))  # 두번쨰 paragraph의 run 갯수
print(doc.paragraphs[1].runs[0].text)  # 두번쨰 paragraph의 첫번째 run 문자열
print(doc.paragraphs[1].runs[1].text)  # 두번쨰 paragraph의 두번째 run 문자열
print(doc.paragraphs[1].runs[2].text)  # 두번쨰 paragraph의 세번째 run 문자열
